#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar Proposta Comercial em formato Word (.docx)
com formatação profissional e elegante
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()

def set_cell_background(cell, fill):
    """Define cor de fundo de célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_header(doc):
    """Adiciona cabeçalho com fundo azul"""
    # Header principal
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title = header.add_run('PROPOSTA COMERCIAL')
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 78, 166)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('Sistema de Gestão de Requisições e Inventário')
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(23, 78, 166)
    
    doc.add_paragraph()
    
    client = doc.add_paragraph()
    client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_run = client.add_run('Igreja Batista Vilas do Atlântico')
    client_run.font.size = Pt(20)
    client_run.font.bold = True
    client_run.font.color.rgb = RGBColor(255, 214, 0)
    
    # Adicionar fundo azul ao cabeçalho
    for para in [header, subtitle, client]:
        para_format = para.paragraph_format
        para_format.space_before = Pt(12)
        para_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run('Data: 03 de Outubro de 2025')
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph('_' * 80)

def add_section_title(doc, icon, title):
    """Adiciona título de seção com ícone"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(12)
    
    run = para.add_run(f'{icon} {title}')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 78, 166)

def add_subsection_title(doc, title):
    """Adiciona subtítulo"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)
    
    run = para.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(18, 58, 123)

def add_info_box(doc, text, bg_color='F0F0F0'):
    """Adiciona caixa de informação com fundo colorido"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Light Grid'
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_background(cell, bg_color)
    
    # Ajustar padding
    cell.paragraphs[0].paragraph_format.space_before = Pt(10)
    cell.paragraphs[0].paragraph_format.space_after = Pt(10)
    cell.paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    cell.paragraphs[0].paragraph_format.right_indent = Inches(0.2)
    
    doc.add_paragraph()

def create_proposta():
    """Cria o documento da proposta"""
    
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ==================== CABEÇALHO ====================
    add_header(doc)
    
    doc.add_paragraph()
    
    # ==================== APRESENTAÇÃO ====================
    add_section_title(doc, '🎯', 'Apresentação do Sistema')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        'Apresento uma solução completa e personalizada de gestão de requisições, eventos e inventário, '
        'desenvolvida especialmente para atender às necessidades operacionais da Igreja Batista Vilas do Atlântico. '
        'O sistema foi criado do zero, utilizando as tecnologias mais modernas do mercado e focando em '
        'proporcionar uma experiência intuitiva e eficiente para todos os níveis de usuários.'
    )
    run.font.size = Pt(11)
    
    # Stats
    doc.add_paragraph()
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    stats = stats_para.add_run('165 HORAS INVESTIDAS  |  3 PLATAFORMAS  |  13 MÓDULOS COMPLETOS')
    stats.font.size = Pt(14)
    stats.font.bold = True
    stats.font.color.rgb = RGBColor(23, 78, 166)
    
    doc.add_paragraph()
    
    # ==================== FUNCIONALIDADES ====================
    add_section_title(doc, '✨', 'Principais Funcionalidades Implementadas')
    
    funcionalidades = [
        ('🔐 Sistema de Autenticação e Segurança', [
            'Sistema JWT profissional com sessões seguras',
            '5 perfis de usuário: Administrador, Pastor, Líder, Secretária, Audiovisual',
            'Controle de acesso granular por funcionalidade',
            'Rastreabilidade total de todas as ações'
        ]),
        ('📅 Gestão Inteligente de Requisições', [
            'Criação rápida e intuitiva de requisições',
            'Detecção automática de conflitos de horário e local',
            'Sugestões inteligentes de horários alternativos',
            'Sistema de prioridades (Normal, Alta, Urgente)',
            'Fluxo de aprovação digital (Pastor/Administrador)',
            'Histórico completo de cada requisição'
        ]),
        ('📦 Controle Avançado de Inventário', [
            'Gestão completa de materiais e equipamentos',
            'Controle de estoque em tempo real',
            'Reserva automática ao aprovar requisições',
            'Alertas de baixo estoque',
            'Categorização por tipo (Áudio, Vídeo, Cabos, Decoração, Esportes)',
            'Histórico de movimentações (entrada/saída/devolução)'
        ]),
        ('🏢 Gestão de Locais e Espaços', [
            'Controle de múltiplos espaços (Anexo 1, Anexo 2, Templo, Salas)',
            'Verificação em tempo real de disponibilidade',
            'Prevenção automática de conflitos de agendamento',
            'Controle de capacidade por espaço'
        ]),
        ('📊 Dashboards Personalizados', [
            'Dashboard Administrativo: Visão geral, estatísticas, filtros avançados',
            'Dashboard Audiovisual: Materiais do dia, retorno de equipamentos',
            'Dashboard do Líder: Suas requisições, status em tempo real'
        ]),
        ('📱 Aplicativo Mobile + PWA', [
            'Versão nativa para iOS e Android',
            'Progressive Web App instalável',
            'Interface otimizada para dispositivos móveis',
            'Notificações push de aprovações e lembretes'
        ]),
        ('🔔 Sistema de Notificações', [
            'Alertas automáticos de conflitos de horário',
            'Notificações de baixo estoque',
            'Avisos de mudança de status de requisições',
            'Lembretes de eventos próximos'
        ])
    ]
    
    for titulo, itens in funcionalidades:
        add_subsection_title(doc, titulo)
        for item in itens:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.size = Pt(10.5)
    
    # ==================== QUEBRA DE PÁGINA ====================
    add_page_break(doc)
    
    # ==================== BENEFÍCIOS ====================
    add_section_title(doc, '📈', 'Benefícios para a IBVA')
    
    beneficios = [
        '90% menos tempo gasto em agendamentos e aprovações manuais',
        '80% de redução no tempo de controle de inventário',
        'Zero conflitos de horário entre eventos',
        'Economia de R$ 5.000/ano evitando perdas e extravios de equipamentos',
        '100% de rastreabilidade em todas as ações do sistema',
        'Acesso em qualquer lugar via web, mobile ou PWA',
        'Comunicação eficiente entre todos os departamentos',
        'Tomada de decisão baseada em dados reais e atualizados'
    ]
    
    for beneficio in beneficios:
        p = doc.add_paragraph(beneficio, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.runs[0]
        run.font.size = Pt(11)
        run.font.bold = True
    
    # ==================== TECNOLOGIAS ====================
    add_section_title(doc, '🛠️', 'Tecnologias Utilizadas')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    tech_text = (
        'Backend: Node.js + Express (usado por Netflix, Uber) | PostgreSQL/Supabase | JWT\n\n'
        'Frontend Web: React (Facebook, Instagram) | Vite | PWA\n\n'
        'Mobile: React Native (Airbnb, Tesla) | Expo | iOS e Android\n\n'
        'Infraestrutura: Vercel (hospedagem premium) | SSL/HTTPS | Backups automáticos | CDN Global'
    )
    
    run = p.add_run(tech_text)
    run.font.size = Pt(10.5)
    
    # ==================== QUEBRA DE PÁGINA ====================
    add_page_break(doc)
    
    # ==================== COMPARAÇÃO DE MERCADO ====================
    add_section_title(doc, '💰', 'Comparação com Valores de Mercado')
    
    # Criar tabela
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Detalhes'
    header_cells[2].text = 'Valor de Mercado'
    
    for cell in header_cells:
        set_cell_background(cell, '174ea6')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
    
    # Dados
    dados = [
        ('Desenvolvimento do Sistema', '165 horas × R$ 200/hora (média mercado)', 'R$ 33.000'),
        ('Aplicativo Mobile', 'App nativo completo iOS e Android', 'R$ 8.000'),
        ('Hospedagem Premium', 'Vercel + Supabase por 12 meses', 'R$ 3.600'),
        ('Suporte Técnico', '12 meses de suporte e manutenção', 'R$ 4.800'),
        ('Treinamento da Equipe', 'Capacitação completa', 'R$ 1.500'),
        ('VALOR TOTAL DE MERCADO', '', 'R$ 50.900')
    ]
    
    for i, (item, detalhe, valor) in enumerate(dados, start=1):
        row = table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = detalhe
        row.cells[2].text = valor
        
        # Última linha em negrito
        if i == len(dados):
            set_cell_background(row.cells[0], 'FFD600')
            set_cell_background(row.cells[1], 'FFD600')
            set_cell_background(row.cells[2], 'FFD600')
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(18, 58, 123)
    
    doc.add_paragraph()
    
    # ==================== MÓDULOS BÔNUS ====================
    add_section_title(doc, '🎁', 'Módulos Bônus Inclusos (Sem Custo Adicional)')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Após a implementação inicial, serão desenvolvidos gratuitamente:')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 78, 166)
    
    add_subsection_title(doc, '💰 Módulo Financeiro [BÔNUS]')
    bonus_financeiro = [
        'Solicitações de recursos financeiros digitalizadas',
        'Fluxo de aprovação financeira',
        'Controle de orçamento por departamento',
        'Relatórios financeiros detalhados',
        'Upload de notas fiscais e comprovantes',
        'Prestação de contas digitalizada'
    ]
    for item in bonus_financeiro:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Valor de Mercado deste módulo: R$ 4.500')
    run.font.bold = True
    run.font.color.rgb = RGBColor(40, 167, 69)
    
    add_subsection_title(doc, '🏥 Módulo Histórico Médico para Missões [BÔNUS]')
    bonus_medico = [
        'Cadastro completo de missionários',
        'Ficha médica detalhada',
        'Histórico de vacinas com alertas de renovação',
        'Medicamentos, alergias e restrições',
        'Contatos de emergência',
        'Relatórios para viagens missionárias'
    ]
    for item in bonus_medico:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Valor de Mercado deste módulo: R$ 4.000')
    run.font.bold = True
    run.font.color.rgb = RGBColor(40, 167, 69)
    
    # ==================== QUEBRA DE PÁGINA ====================
    add_page_break(doc)
    
    # ==================== NOTA PESSOAL ====================
    add_section_title(doc, '💙', 'Uma Nota Pessoal')
    
    nota_table = doc.add_table(rows=1, cols=1)
    cell = nota_table.rows[0].cells[0]
    set_cell_background(cell, 'FFF3CD')
    
    nota_text = (
        'Como este sistema foi desenvolvido para a minha igreja, onde congrego e sirvo, '
        'não posso encarar este projeto apenas como uma transação comercial. Esta é minha forma de contribuir '
        'com a obra de Deus através dos talentos que Ele me deu. Por isso, o valor cobrado reflete meu '
        'compromisso espiritual e não o valor de mercado do trabalho realizado.\n\n'
        'Embora o valor de mercado deste sistema completo seja de R$ 50.900, meu desejo '
        'é torná-lo acessível para que a IBVA possa se beneficiar desta tecnologia sem comprometer o orçamento '
        'da igreja para outras áreas importantes do ministério.\n\n'
        'Este não é apenas um software, mas uma ferramenta que facilitará o trabalho de todos os '
        'departamentos, permitindo que a equipe dedique mais tempo ao que realmente importa: '
        'cuidar das pessoas e servir ao Reino de Deus.'
    )
    
    para = cell.paragraphs[0]
    para.text = nota_text
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(133, 100, 4)
    
    doc.add_paragraph()
    
    # ==================== INVESTIMENTO ====================
    add_section_title(doc, '💎', 'Investimento Proposto')
    
    # Caixa de preço
    price_table = doc.add_table(rows=1, cols=1)
    price_cell = price_table.rows[0].cells[0]
    set_cell_background(price_cell, 'FFD600')
    
    price_para = price_cell.paragraphs[0]
    price_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    price_para.paragraph_format.space_before = Pt(20)
    price_para.paragraph_format.space_after = Pt(20)
    
    # Título
    run1 = price_para.add_run('VALOR DO SISTEMA COMPLETO\n\n')
    run1.font.size = Pt(14)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(18, 58, 123)
    
    # Preço principal
    run2 = price_para.add_run('R$ 3.000\n\n')
    run2.font.size = Pt(36)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(18, 58, 123)
    
    # Descrição
    run3 = price_para.add_run('Pagamento único\n')
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(18, 58, 123)
    
    run4 = price_para.add_run('Inclui: Sistema completo + App Mobile + 2 Módulos Bônus + Treinamento\n\n')
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(18, 58, 123)
    
    # Linha divisória
    run5 = price_para.add_run('_' * 50 + '\n\n')
    run5.font.color.rgb = RGBColor(18, 58, 123)
    
    # Manutenção
    run6 = price_para.add_run('Manutenção: R$ 300/mês\n')
    run6.font.size = Pt(16)
    run6.font.bold = True
    run6.font.color.rgb = RGBColor(18, 58, 123)
    
    run7 = price_para.add_run('Período: 12 meses | Total: R$ 3.600/ano')
    run7.font.size = Pt(11)
    run7.font.color.rgb = RGBColor(18, 58, 123)
    
    doc.add_paragraph()
    
    # Caixa de total
    total_table = doc.add_table(rows=1, cols=1)
    total_cell = total_table.rows[0].cells[0]
    set_cell_background(total_cell, 'E7F3FF')
    
    total_para = total_cell.paragraphs[0]
    total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_para.paragraph_format.space_before = Pt(15)
    total_para.paragraph_format.space_after = Pt(15)
    
    run1 = total_para.add_run('INVESTIMENTO TOTAL NO PRIMEIRO ANO\n\n')
    run1.font.size = Pt(14)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(23, 78, 166)
    
    run2 = total_para.add_run('R$ 6.600\n\n')
    run2.font.size = Pt(28)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(18, 58, 123)
    
    run3 = total_para.add_run('(R$ 3.000 sistema + R$ 3.600 manutenção anual)\n\n')
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    
    run4 = total_para.add_run('Economia de 87% sobre o valor de mercado (R$ 50.900)')
    run4.font.size = Pt(12)
    run4.font.bold = True
    run4.font.color.rgb = RGBColor(40, 167, 69)
    
    doc.add_paragraph()
    
    # ==================== QUEBRA DE PÁGINA ====================
    add_page_break(doc)
    
    # ==================== O QUE ESTÁ INCLUÍDO ====================
    add_section_title(doc, '✅', 'O que está Incluído no Investimento')
    
    inclusos = [
        'Sistema Web completo e responsivo em produção',
        'Aplicativo Mobile nativo (iOS e Android)',
        'Progressive Web App (instalável)',
        '13 módulos funcionais completos',
        'Módulo Financeiro (bônus - desenvolvimento em 30 dias)',
        'Módulo Histórico Médico (bônus - desenvolvimento em 30 dias)',
        'Hospedagem premium Vercel por 12 meses',
        'Banco de dados Supabase por 12 meses',
        'Certificado SSL/HTTPS incluso',
        'Backups automáticos diários',
        'Suporte técnico por 12 meses via WhatsApp/Email',
        'Atualizações de segurança e correções de bugs',
        'Treinamento completo da equipe',
        'Documentação técnica e manual do usuário',
        'Garantia de funcionamento'
    ]
    
    for item in inclusos:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.size = Pt(10.5)
    
    # ==================== CRONOGRAMA ====================
    add_section_title(doc, '📅', 'Cronograma de Implementação')
    
    crono_table = doc.add_table(rows=6, cols=3)
    crono_table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    header_cells = crono_table.rows[0].cells
    header_cells[0].text = 'Fase'
    header_cells[1].text = 'Atividade'
    header_cells[2].text = 'Prazo'
    
    for cell in header_cells:
        set_cell_background(cell, '174ea6')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Dados do cronograma
    crono_dados = [
        ('Imediato', 'Sistema completo já disponível e funcionando', '✅ Concluído'),
        ('Semana 1', 'Treinamento da equipe e ajustes iniciais', '5 dias úteis'),
        ('Mês 1', 'Desenvolvimento completo do Módulo Financeiro', '30 dias'),
        ('Mês 2', 'Desenvolvimento do Módulo Histórico Médico', '30 dias'),
        ('Meses 3-12', 'Suporte, manutenção e melhorias contínuas', 'Contínuo')
    ]
    
    for i, (fase, atividade, prazo) in enumerate(crono_dados, start=1):
        row = crono_table.rows[i]
        row.cells[0].text = fase
        row.cells[1].text = atividade
        row.cells[2].text = prazo
    
    doc.add_paragraph()
    
    # ==================== FORMAS DE PAGAMENTO ====================
    add_section_title(doc, '💳', 'Formas de Pagamento')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pagamento_text = (
        'Sistema (R$ 3.000): PIX, transferência bancária ou depósito\n\n'
        'Manutenção Mensal (R$ 300/mês): Pagamento mensal via PIX ou transferência\n\n'
        'Primeira cobrança de manutenção: 30 dias após assinatura do contrato\n\n'
        'Vencimento: Todo dia 10 de cada mês'
    )
    
    run = p.add_run(pagamento_text)
    run.font.size = Pt(11)
    
    # ==================== QUEBRA DE PÁGINA ====================
    add_page_break(doc)
    
    # ==================== AGRADECIMENTO ====================
    add_section_title(doc, '🙏', 'Agradecimento Final')
    
    agradecimento_table = doc.add_table(rows=1, cols=1)
    agr_cell = agradecimento_table.rows[0].cells[0]
    set_cell_background(agr_cell, '174ea6')
    
    agr_text = (
        'Agradeço imensamente a oportunidade de contribuir com a Igreja Batista Vilas do Atlântico através '
        'deste sistema. É uma honra poder usar os talentos que Deus me deu para facilitar a gestão e '
        'administração da nossa igreja, permitindo que a liderança possa dedicar mais tempo ao que '
        'realmente importa: cuidar das pessoas e expandir o Reino de Deus.\n\n'
        'Desenvolvi este sistema com muito carinho, dedicação e oração, pensando em cada detalhe para '
        'que ele seja realmente útil e transformador para todos os ministérios da igreja.\n\n'
        'Estou à disposição para esclarecer qualquer dúvida e ansioso para ver este sistema '
        'trazendo benefícios práticos para todos os departamentos da IBVA.\n\n'
        'Que Deus abençoe ricamente a Igreja Batista Vilas do Atlântico e todo o seu ministério! 🙏'
    )
    
    agr_para = agr_cell.paragraphs[0]
    agr_para.text = agr_text
    agr_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    agr_para.paragraph_format.space_before = Pt(15)
    agr_para.paragraph_format.space_after = Pt(15)
    agr_para.paragraph_format.left_indent = Inches(0.2)
    agr_para.paragraph_format.right_indent = Inches(0.2)
    
    for run in agr_para.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    # ==================== CONTATO ====================
    add_section_title(doc, '📞', 'Contato para Dúvidas e Esclarecimentos')
    
    contato_para = doc.add_paragraph()
    contato_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contato_text = (
        'Desenvolvedor Responsável: Maurício Oliveira\n'
        '📧 Email: mauriciooliveira@exemplo.com\n'
        '📱 WhatsApp: (71) 9xxxx-xxxx\n'
        '⏰ Horário de Atendimento: Segunda a Sexta, 9h às 18h'
    )
    
    run = contato_para.add_run(contato_text)
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    # Rodapé
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Proposta Comercial - Sistema de Gestão IBVA | Versão 1.0 | Outubro 2025')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # ==================== SALVAR ====================
    filename = '/Users/mauriciooliveira/Desktop/sistemaderequerimento/PROPOSTA_COMERCIAL_IBVA.docx'
    doc.save(filename)
    print(f'✅ Proposta criada com sucesso!')
    print(f'📄 Arquivo salvo em: {filename}')
    print(f'📊 Tamanho aproximado: {len(doc.paragraphs)} parágrafos')
    print(f'🎨 Formatação: Profissional e elegante')
    print(f'💰 Valor proposto: R$ 3.000 + R$ 300/mês')

if __name__ == '__main__':
    try:
        create_proposta()
    except ImportError:
        print('❌ Erro: Biblioteca python-docx não instalada')
        print('📦 Instale com: pip install python-docx')
    except Exception as e:
        print(f'❌ Erro ao criar proposta: {str(e)}')


